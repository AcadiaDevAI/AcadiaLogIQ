import json
import logging
import hashlib
import sys
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path 
from typing import Dict, List, Optional, Tuple, TypedDict, Set

import boto3
from botocore.config import Config as BotoConfig
from fastapi import (
    BackgroundTasks, Depends, FastAPI, File, Header,
    HTTPException, Query, Request, UploadFile, status,
)
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address
from pydantic import BaseModel, ConfigDict, Field, field_validator

from backend.config import settings
from backend.clerk_auth import clerk_auth_dependency, is_clerk_enabled, get_clerk_user_display
from backend.storage.local_storage import LocalStorageProvider
from backend.vector_store import (
    count_pg_chunks,
    create_ingestion_job,
    delete_all_chat_sessions,
    delete_chat_session,
    delete_document_and_chunks,
    get_bm25_index,
    get_chat_session,
    get_ingestion_job,
    insert_document_and_chunks,
    list_active_files,
    list_chat_sessions,
    pgvector_search,
    purge_orphan_chunks_db,
    rebuild_bm25_from_postgres,
    reset_pg_data,
    save_message_to_session,
    update_ingestion_job,
    update_message_feedback,
)
from backend.services.contextual_ingestion_service import process_document
from backend.vector_store import find_duplicate_by_hash, find_version_candidates
from backend.retrieval.orchestrator import retrieve as orchestrator_retrieve




if sys.version_info < (3, 11):
    raise RuntimeError("Python 3.11+ required")


logging.basicConfig(
    level=getattr(logging, settings.LOG_LEVEL.upper(), logging.INFO),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("acadia-log-iq")


CHARS_PER_TOKEN = 3

TOKEN_BUDGET = {
    "MODEL_MAX_TOKENS": 32_768,
    "MAX_LOG_CONTEXT_CHARS": 24_000,
    "MAX_KB_CONTEXT_CHARS": 24_000,
    "MAX_GENERATION_TOKENS": 2_048,
    "MAX_SINGLE_CHUNK_CHARS": 6_000,
    "PROMPT_OVERHEAD_CHARS": 1_500,
}

_MAX_PROMPT_TOKENS = TOKEN_BUDGET["MODEL_MAX_TOKENS"] - TOKEN_BUDGET["MAX_GENERATION_TOKENS"]
MAX_TOTAL_PROMPT_CHARS = (_MAX_PROMPT_TOKENS * CHARS_PER_TOKEN) - TOKEN_BUDGET["PROMPT_OVERHEAD_CHARS"]

CHUNK_CONFIG = {
    "LINES_PER_CHUNK": 80,
    "MAX_CHUNK_CHARS": 6_000,
    "OVERLAP_LINES": 10,
}

HYBRID_CONFIG = {
    "VECTOR_WEIGHT": 0.6,
    "BM25_WEIGHT": 0.4,
    "VECTOR_CANDIDATES": 20,
    "BM25_CANDIDATES": 20,
    "RERANK_TOP_K_LOG": 6,
    "RERANK_TOP_K_KB": 5,
}

storage = LocalStorageProvider()


class JobInfo(TypedDict, total=False):
    job_id: str
    status: str
    processed_chunks: int
    total_chunks: int
    successful_chunks: int
    file: Optional[str]
    file_type: Optional[str]
    file_size_mb: float
    file_hash: str
    error: Optional[str]
    created_at: datetime
    completed_at: Optional[datetime]
    owner_id: Optional[str]
    file_id: Optional[str]


bm25 = None


def _make_bedrock_client():
    boto_cfg = BotoConfig(
        retries={"max_attempts": 10, "mode": "adaptive"},
        read_timeout=120, connect_timeout=30, tcp_keepalive=True,
    )
    kwargs = {
        "service_name": "bedrock-runtime",
        "region_name": settings.AWS_REGION,
        "config": boto_cfg,
    }
    if settings.AWS_ACCESS_KEY_ID and settings.AWS_SECRET_ACCESS_KEY:
        kwargs["aws_access_key_id"] = settings.AWS_ACCESS_KEY_ID
        kwargs["aws_secret_access_key"] = settings.AWS_SECRET_ACCESS_KEY
        if settings.AWS_SESSION_TOKEN:
            kwargs["aws_session_token"] = settings.AWS_SESSION_TOKEN
    return boto3.client(**kwargs)


bedrock = _make_bedrock_client()


def _make_ses_client():
    ses_region = settings.SES_REGION or settings.AWS_REGION
    kwargs = {
        "service_name": "ses",
        "region_name": ses_region,
    }
    if settings.AWS_ACCESS_KEY_ID and settings.AWS_SECRET_ACCESS_KEY:
        kwargs["aws_access_key_id"] = settings.AWS_ACCESS_KEY_ID
        kwargs["aws_secret_access_key"] = settings.AWS_SECRET_ACCESS_KEY
        if settings.AWS_SESSION_TOKEN:
            kwargs["aws_session_token"] = settings.AWS_SESSION_TOKEN
    return boto3.client(**kwargs)


ses_client = None
try:
    if settings.SES_ENABLED.lower() == "true":
        ses_client = _make_ses_client()
        logger.info("SES ready")
except Exception as e:
    logger.warning("SES init failed: %s", e)


def send_feedback_email(subject: str, body_text: str, body_html: str) -> bool:
    if not ses_client:
        return False
    try:
        ses_client.send_email(
            Source=settings.SES_SENDER_EMAIL,
            Destination={"ToAddresses": [settings.SES_FEEDBACK_RECIPIENT]},
            Message={
                "Subject": {"Data": subject, "Charset": "UTF-8"},
                "Body": {
                    "Text": {"Data": body_text, "Charset": "UTF-8"},
                    "Html": {"Data": body_html, "Charset": "UTF-8"},
                },
            },
        )
        return True
    except Exception as e:
        logger.exception("SES send_email failed: %s", e)
        return False


def _resolve_user_display(user_id: Optional[str]) -> dict:
    return get_clerk_user_display(user_id)


def _normalize_owner_id(user_id: Optional[str]) -> str:
    return user_id or "anonymous"


def _get_active_indexed_file_ids(user_id: Optional[str]) -> Set[str]:
    owner_id = _normalize_owner_id(user_id)
    return {f["id"] for f in list_active_files(owner_id) if f.get("status") == "indexed"}


@asynccontextmanager
async def lifespan(app: FastAPI):
    global bm25
    logger.info("Starting API...")
    bm25 = get_bm25_index()
    doc_count = rebuild_bm25_from_postgres()
    logger.info("BM25 ready from PostgreSQL: %d docs", doc_count)
    yield
    logger.info("Shutting down...")


app = FastAPI(
    title="Acadia's Log IQ API",
    description="AI log analysis — Hybrid Search + Re-ranking",
    version="2.2.0",
    lifespan=lifespan,
)

limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
         "http://18.233.93.19:8501",
        "http://localhost:8501",
        "http://127.0.0.1:8501",
        "http://localhost:8001",
        "http://127.0.0.1:8001",
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "http://localhost:8000",
        "http://127.0.0.1:8000",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Processing-Time"],
    max_age=600,
)


def verify_api_key(x_api_key: Optional[str] = Header(default=None, alias="X-API-Key")) -> bool:
    if settings.API_KEY:
        if not x_api_key or x_api_key != settings.API_KEY:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid API key")
    return True


async def auth_dependency(
    request: Request,
    x_api_key: Optional[str] = Header(default=None, alias="X-API-Key"),
) -> Optional[str]:
    if is_clerk_enabled():
        return await clerk_auth_dependency(request)
    verify_api_key(x_api_key)
    return None


class Question(BaseModel):
    q: str = Field(min_length=1, max_length=1000)
    session_id: Optional[str] = None
    model_config = ConfigDict(extra="ignore")

    @field_validator("q")
    @classmethod
    def validate_q(cls, v):
        v = (v or "").strip()
        if not v:
            raise ValueError("Empty")
        return v


class UploadResponse(BaseModel):
    job_id: str
    file_id: str
    message: str
    file_hash: Optional[str] = None
    model_config = ConfigDict(extra="ignore")


class JobStatus(BaseModel):
    job_id: str
    status: str
    processed_chunks: int = 0
    total_chunks: Optional[int] = None
    successful_chunks: Optional[int] = None
    file: Optional[str] = None
    file_type: Optional[str] = None
    file_size_mb: Optional[float] = None
    file_hash: Optional[str] = None
    error: Optional[str] = None
    created_at: datetime
    completed_at: Optional[datetime] = None
    model_config = ConfigDict(extra="ignore")


class AnswerResponse(BaseModel):
    answer: str
    sources: List[str]
    confidence: float = Field(ge=0, le=1)
    processing_time_ms: Optional[int] = None
    context_stats: Optional[Dict] = None
    session_id: Optional[str] = None
    model_config = ConfigDict(extra="ignore")


class ChatMessage(BaseModel):
    role: str
    content: str
    sources: Optional[Dict] = None
    timestamp: str
    feedback: Optional[str] = None


class ChatSession(BaseModel):
    id: str
    title: str
    messages: List[ChatMessage]
    created_at: str
    updated_at: str


class FileInfo(BaseModel):
    id: str
    name: str
    file_type: str
    size_mb: float
    status: str
    job_id: Optional[str] = None
    uploaded_at: str
    owner_id: Optional[str] = None


def safe_embed(text: str) -> Optional[List[float]]:
    if not text or not text.strip():
        return None
    try:
        body = json.dumps({"inputText": text[: settings.MAX_CHARS]}).encode("utf-8")
        resp = bedrock.invoke_model(
            modelId=settings.BEDROCK_EMBED_MODEL,
            body=body,
            accept="application/json",
            contentType="application/json",
        )
        payload = json.loads(resp["body"].read().decode("utf-8"))
        emb = payload.get("embedding")
        return emb if isinstance(emb, list) else None
    except Exception as e:
        logger.exception("Embed failed: %s", e)
        return None


def estimate_tokens(text: str) -> int:
    return len(text) // CHARS_PER_TOKEN


def safe_generate(prompt: str, max_tokens: int = None) -> str:
    if max_tokens is None:
        max_tokens = TOKEN_BUDGET["MAX_GENERATION_TOKENS"]
    try:
        max_prompt_tokens = TOKEN_BUDGET["MODEL_MAX_TOKENS"] - max_tokens - 200
        max_prompt_chars = max_prompt_tokens * CHARS_PER_TOKEN

        if len(prompt) > max_prompt_chars:
            logger.warning("TRUNCATING: %d -> %d chars", len(prompt), max_prompt_chars)
            marker = "\nANSWER:"
            pos = prompt.rfind(marker)
            if pos > 0:
                tail = prompt[pos:]
                prompt = (
                    prompt[: max_prompt_chars - len(tail) - 80]
                    + "\n\n[... context truncated ...]\n"
                    + tail
                )
            else:
                prompt = prompt[:max_prompt_chars] + "\n\n[... truncated ...]\n"

        body = json.dumps(
            {
                "prompt": prompt,
                "max_tokens": max_tokens,
                "temperature": 0.1,
                "top_p": 0.9,
            }
        ).encode("utf-8")

        resp = bedrock.invoke_model(
            modelId=settings.BEDROCK_LLM_MODEL,
            body=body,
            accept="application/json",
            contentType="application/json",
        )
        payload = json.loads(resp["body"].read().decode("utf-8"))

        if isinstance(payload, dict):
            if "outputs" in payload and payload["outputs"]:
                return (payload["outputs"][0].get("text") or "").strip() or "No response."
            if "generation" in payload:
                return str(payload["generation"]).strip()
            if "outputText" in payload:
                return str(payload["outputText"]).strip()
        return "No response generated."
    except Exception as e:
        logger.exception("Generation failed: %s", e)
        return "Error generating response. Please try again."


def hybrid_search(
    query: str,
    query_embedding: List[float],
    file_type: str,
    n_results: int = 10,
    allowed_file_ids: Optional[Set[str]] = None,
    owner_id: Optional[str] = None,
) -> List[Tuple[str, str, Dict, float]]:
    v_weight = HYBRID_CONFIG["VECTOR_WEIGHT"]
    b_weight = HYBRID_CONFIG["BM25_WEIGHT"]
    rrf_k = 60
    owner_id = _normalize_owner_id(owner_id)

    vector_results = {}
    try:
        vector_hits = pgvector_search(
            query_embedding=query_embedding,
            n_results=HYBRID_CONFIG["VECTOR_CANDIDATES"],
            allowed_file_ids=list(allowed_file_ids) if allowed_file_ids else None,
        )

        rank = 0
        for hit in vector_hits:
            doc_id = hit["id"]
            meta = hit["metadata"]
            meta_file_id = meta.get("file_id")
            meta_owner_id = meta.get("owner_id", "anonymous")

            if allowed_file_ids is not None and meta_file_id not in allowed_file_ids:
                continue
            if meta_owner_id != owner_id:
                continue

            rank += 1
            vector_results[doc_id] = {
                "text": hit["text"],
                "metadata": meta,
                "rank": rank,
                "similarity": max(0.0, 1.0 - float(hit["distance"])),
            }

    except Exception as e:
        logger.warning("Vector search failed (%s): %s", file_type, e)

    bm25_results = {}
    try:
        if bm25 and bm25.size > 0:
            raw_bm25 = bm25.search(
                query,
                n_results=HYBRID_CONFIG["BM25_CANDIDATES"],
                file_type=file_type,
            )
            rank = 0
            for doc_id, text_value, meta, score in raw_bm25:
                meta_file_id = meta.get("file_id")
                meta_owner_id = meta.get("owner_id", "anonymous")

                if allowed_file_ids is not None and meta_file_id not in allowed_file_ids:
                    continue
                if meta_owner_id != owner_id:
                    continue

                rank += 1
                bm25_results[doc_id] = {
                    "text": text_value,
                    "metadata": meta,
                    "rank": rank,
                    "bm25_score": score,
                }

                if rank >= HYBRID_CONFIG["BM25_CANDIDATES"]:
                    break

    except Exception as e:
        logger.warning("BM25 search failed: %s", e)

    combined: Dict[str, float] = {}
    all_data: Dict[str, dict] = {}

    for doc_id, data in vector_results.items():
        combined[doc_id] = combined.get(doc_id, 0) + v_weight / (rrf_k + data["rank"])
        all_data[doc_id] = data

    for doc_id, data in bm25_results.items():
        combined[doc_id] = combined.get(doc_id, 0) + b_weight / (rrf_k + data["rank"])
        if doc_id not in all_data:
            all_data[doc_id] = data

    sorted_ids = sorted(combined, key=lambda x: combined[x], reverse=True)

    return [
        (did, all_data[did]["text"], all_data[did]["metadata"], combined[did])
        for did in sorted_ids[:n_results]
    ]


def rerank_chunks(
    query: str,
    chunks: List[Tuple[str, str, Dict, float]],
    top_k: int = 6,
) -> List[Tuple[str, str, Dict, float]]:
    if not chunks or len(chunks) <= 1:
        return chunks[:top_k]

    candidates = chunks[: min(len(chunks), 12)]

    previews = []
    for i, (_, text, meta, _) in enumerate(candidates):
        preview = text[:600].replace("\n", " ").strip()
        src = meta.get("source", "?")
        previews.append(f"[{i + 1}] (source: {src}) {preview}")

    rerank_prompt = f"""Rate each chunk's relevance to the question (0=irrelevant, 10=perfect match).
Question: {query}

Chunks:
{chr(10).join(previews)}

Respond ONLY with a JSON array: [{{"chunk":1,"score":8}}, ...]"""

    try:
        resp = safe_generate(rerank_prompt, max_tokens=512)
        raw = resp.strip()
        if "```" in raw:
            raw = raw.split("```")[1] if "```json" not in raw else raw.split("```json")[1].split("```")[0]
        start, end = raw.find("["), raw.rfind("]") + 1
        if start < 0 or end <= start:
            raise ValueError("No JSON array found")

        scores = json.loads(raw[start:end])

        scored = []
        for item in scores:
            idx = item.get("chunk", 0) - 1
            relevance = item.get("score", 0)
            if 0 <= idx < len(candidates):
                did, text, meta, orig = candidates[idx]
                final = (relevance / 10.0) * 0.7 + orig * 30 * 0.3
                scored.append((did, text, meta, final))

        scored_ids = {s[0] for s in scored}
        for c in candidates:
            if c[0] not in scored_ids:
                scored.append(c)

        scored.sort(key=lambda x: x[3], reverse=True)
        return scored[:top_k]

    except Exception as e:
        logger.warning("Re-rank failed (%s), using hybrid scores", e)
        return candidates[:top_k]


def truncate_chunk(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    t = text[:max_chars]
    nl = t.rfind("\n")
    if nl > max_chars * 0.7:
        t = t[:nl]
    return t + "\n[... truncated ...]"


def assemble_context(
    ranked: List[Tuple[str, str, Dict, float]],
    max_total_chars: int,
    max_sources: int = 5,
) -> Tuple[str, List[str]]:
    max_chunk = TOKEN_BUDGET["MAX_SINGLE_CHUNK_CHARS"]
    if not ranked:
        return "", []

    parts = []
    source_scores: Dict[str, float] = {}
    total = 0

    for _, text, meta, score in ranked:
        if not text or not text.strip():
            continue

        chunk = truncate_chunk(text.strip(), max_chunk)
        src = meta.get("source", "unknown")
        entry = f"[Source: {src}]\n{chunk}"

        if total + len(entry) > max_total_chars:
            remaining = max_total_chars - total
            if remaining > 300:
                parts.append(f"[Source: {src}]\n{truncate_chunk(chunk, remaining - 60)}")
                source_scores[src] = source_scores.get(src, 0.0) + score
            break

        parts.append(entry)
        source_scores[src] = source_scores.get(src, 0.0) + score
        total += len(entry)

    if not source_scores:
        return "\n\n".join(parts), []

    ranked_sources = sorted(source_scores.items(), key=lambda x: x[1], reverse=True)
    top_score = ranked_sources[0][1]
    threshold = top_score * 0.40
    final_sources = [src for src, score in ranked_sources if score >= threshold]
    return "\n\n".join(parts), final_sources[:max_sources]

def has_sufficient_document_support(
    question: str,
    ranked: List[Tuple[str, str, Dict, float]],
    min_score: float = 0.18,
    min_keyword_hits: int = 1,
) -> bool:
    """
    Hard gate:
    only answer when retrieval shows actual document support.

    Rules:
    - require at least one reranked chunk
    - require top score above threshold
    - require at least one meaningful keyword overlap between question and retrieved text
    """
    if not ranked:
        return False

    top_score = float(ranked[0][3] or 0.0)
    if top_score < min_score:
        return False

    import re

    stop_words = {
        "the", "a", "an", "is", "are", "to", "for", "of", "in", "on", "how",
        "what", "when", "where", "why", "do", "does", "can", "i", "me", "my",
        "you", "your", "please", "tell", "about"
    }

    q_terms = {
        t for t in re.findall(r"\w+", (question or "").lower())
        if len(t) > 2 and t not in stop_words
    }
    if not q_terms:
        return True

    combined_text = " ".join((item[1] or "") for item in ranked[:3]).lower()
    hit_count = sum(1 for term in q_terms if term in combined_text)

    return hit_count >= min_keyword_hits


def calculate_file_hash_bytes(content: bytes) -> str:
    sha = hashlib.sha256()
    sha.update(content)
    return sha.hexdigest()


def extract_text_from_pdf(fp: Path) -> str:
    try:
        import fitz

        parts = []
        with fitz.open(str(fp)) as doc:
            for i, page in enumerate(doc):
                t = page.get_text("text")
                if t and t.strip():
                    parts.append(f"--- Page {i+1} ---\n{t}")
        return "\n\n".join(parts) if parts else ""
    except ImportError:
        pass

    try:
        import pdfplumber

        parts = []
        with pdfplumber.open(str(fp)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip():
                    parts.append(f"--- Page {i+1} ---\n{t}")
        return "\n\n".join(parts) if parts else ""
    except ImportError:
        raise RuntimeError("No PDF library. pip install PyMuPDF")


def extract_text_from_docx(fp: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("pip install python-docx")

    doc = Document(str(fp))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for ti, table in enumerate(doc.tables):
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
        if rows:
            parts.append(f"--- Table {ti + 1} ---")
            parts.extend(rows)
    return "\n".join(parts)


def extract_text(fp: Path) -> str:
    ext = fp.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(fp)
    if ext == ".docx":
        return extract_text_from_docx(fp)
    return fp.read_text(encoding="utf-8", errors="ignore")


def iter_text_chunks(text, max_chars=None, lines_per=None, overlap=None):
    max_chars = max_chars or CHUNK_CONFIG["MAX_CHUNK_CHARS"]
    lines_per = lines_per or CHUNK_CONFIG["LINES_PER_CHUNK"]
    overlap = overlap or CHUNK_CONFIG["OVERLAP_LINES"]
    lines = text.splitlines(keepends=True)
    if not lines:
        return
    buf, sz = [], 0
    for ln in lines:
        buf.append(ln)
        sz += len(ln)
        if len(buf) >= lines_per or sz >= max_chars:
            yield "".join(buf)
            if overlap > 0 and len(buf) > overlap:
                buf = buf[-overlap:]
                sz = sum(len(l) for l in buf)
            else:
                buf.clear()
                sz = 0
    if buf:
        yield "".join(buf)


def iter_line_chunks(fp: Path, lines_per=None):
    lines_per = lines_per or CHUNK_CONFIG["LINES_PER_CHUNK"]
    ext = fp.suffix.lower()
    if ext in (".pdf", ".docx"):
        text = extract_text(fp)
        if not text or not text.strip():
            return
        for c in iter_text_chunks(text, lines_per=lines_per):
            yield c
        return

    text = fp.read_text(encoding="utf-8", errors="ignore")
    if text and text.strip():
        for c in iter_text_chunks(text, lines_per=lines_per):
            yield c


async def index_file_job(
    job_id: str,
    storage_uri: str,
    filename: str,
    file_type: str,
    file_id: str,
    owner_id: Optional[str],
    file_size_mb: float,
):
    owner_id = _normalize_owner_id(owner_id)
    update_ingestion_job(job_id, status="running")

    try:
        local_path = storage.resolve_local_path(storage_uri)
        if not local_path or not local_path.exists():
            raise RuntimeError(f"Stored file path is not readable: {storage_uri}")

        job = get_ingestion_job(job_id)
        file_hash = job["file_hash"] if job else ""

        processed = process_document(
            local_path=local_path,
            filename=filename,
            file_type=file_type,
            owner_id=owner_id,
            fingerprint=file_hash,
            exact_duplicate_lookup=find_duplicate_by_hash,
            version_candidate_lookup=find_version_candidates,
        )

        if processed["status"] == "exact_duplicate":
            update_ingestion_job(
                job_id,
                status="done",
                processed_chunks="0",
                total_chunks="0",
                successful_chunks="0",
                error=None,
                completed_at=datetime.now(timezone.utc),
            )
            return

        chunk_rows = []
        bm25_ids, bm25_docs, bm25_metas = [], [], []

        total_chunks = len(processed["chunk_rows"])

        # --- Concurrent embedding ---
        embed_inputs: List[Tuple[int, Dict[str, Any], str]] = []
        for idx, row in enumerate(processed["chunk_rows"]):
            embed_text = row.get("contextualized_content") or row["content"]
            embed_inputs.append((idx, row, embed_text))

        embeddings_map: Dict[int, List[float]] = {}
        embed_workers = min(settings.EMBED_CONCURRENCY, max(1, total_chunks))

        with ThreadPoolExecutor(max_workers=embed_workers) as executor:
            future_to_idx = {
                executor.submit(safe_embed, text): idx
                for idx, _row, text in embed_inputs
            }
            for future in as_completed(future_to_idx):
                idx = future_to_idx[future]
                try:
                    emb = future.result()
                    if emb:
                        embeddings_map[idx] = emb
                except Exception:
                    pass

        for idx, row, embed_text in embed_inputs:
            emb = embeddings_map.get(idx)
            if not emb:
                continue

            chunk_id = f"{file_id}:{job_id}:{idx}"
            row["id"] = chunk_id
            row["embedding"] = emb
            chunk_rows.append(row)

            bm25_ids.append(chunk_id)
            bm25_docs.append(embed_text)
            bm25_metas.append(
                {
                    "file_id": file_id,
                    "owner_id": owner_id,
                    "source": filename,
                    "file_type": file_type,
                    "section_heading": row.get("section_heading"),
                    "chunk_type": row.get("chunk_type"),
                    "summary": row.get("summary"),
                    "labels_json": row.get("labels_json", {}),
                    "metadata_json": row.get("metadata_json", {}),
                }
            )

            if (len(chunk_rows)) % settings.BATCH_SIZE == 0:
                update_ingestion_job(
                    job_id,
                    processed_chunks=str(len(chunk_rows)),
                    total_chunks=str(total_chunks),
                    successful_chunks=str(len(chunk_rows)),
                )

        inserted = insert_document_and_chunks(
            document_id=file_id,
            filename=filename,
            fingerprint=file_hash,
            chunk_rows=chunk_rows,
            owner_id=owner_id,
            file_type=file_type,
            storage_uri=storage_uri,
            file_size_mb=file_size_mb,
            metadata=processed["document_metadata"],
            version_decision=processed["version_decision"],
        )

        if bm25 and inserted["status"] == "inserted" and bm25_ids:
            bm25.remove_documents_by_file_id(file_id)
            bm25.add_documents_batch(bm25_ids, bm25_docs, bm25_metas)

        update_ingestion_job(
            job_id,
            status="done",
            processed_chunks=str(total_chunks),
            total_chunks=str(total_chunks),
            successful_chunks=str(len(chunk_rows)),
            error=None,
            completed_at=datetime.now(timezone.utc),
        )

    except Exception as exc:
        logger.exception("Indexing failed for %s: %s", filename, exc)
        update_ingestion_job(
            job_id,
            status="error",
            error=str(exc),
            completed_at=datetime.now(timezone.utc),
        )   


@app.middleware("http")
async def log_requests(request: Request, call_next):
    import time

    start = time.perf_counter()
    response = await call_next(request)
    ms = (time.perf_counter() - start) * 1000.0
    response.headers["X-Processing-Time"] = f"{ms:.2f}ms"
    logger.info("%s %s -> %s (%.2fms)", request.method, request.url.path, response.status_code, ms)
    return response


@app.get("/health")
async def health_check():
    chunk_count = 0
    try:
        chunk_count = count_pg_chunks()
    except Exception:
        pass

    return {
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "model": settings.BEDROCK_LLM_MODEL,
        "services": {
            "vector_store": f"{chunk_count} chunks" if chunk_count >= 0 else "uninitialized",
            "bm25_index": f"{bm25.size} docs" if bm25 else "uninitialized",
            "bedrock": "available",
        },
        "search_mode": "hybrid (pgvector + BM25 + re-ranking)",
        "auth_mode": "clerk" if is_clerk_enabled() else ("api_key" if settings.API_KEY else "open"),
    }


@app.get("/me")
async def get_current_user(request: Request, user_id: Optional[str] = Depends(auth_dependency)):
    if is_clerk_enabled() and user_id:
        payload = getattr(request.state, "clerk_payload", {})
        return {
            "authenticated": True,
            "user_id": user_id,
            "issuer": payload.get("iss"),
            "auth_mode": "clerk",
        }
    return {
        "authenticated": False,
        "user_id": None,
        "auth_mode": "api_key" if settings.API_KEY else "open",
    }


@app.post("/reset")
async def reset_all(user_id: Optional[str] = Depends(auth_dependency)):
    deleted_chunks = reset_pg_data()
    if bm25:
        bm25.clear()

    return {
        "status": "success",
        "message": "All data deleted",
        "deleted_chunks": deleted_chunks,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.post("/purge_orphans")
async def purge_orphan_chunks(user_id: Optional[str] = Depends(auth_dependency)):
    deleted = purge_orphan_chunks_db()
    bm25_count = rebuild_bm25_from_postgres()

    return {
        "status": "purged",
        "orphans_deleted": deleted,
        "chunks_remaining": count_pg_chunks(),
        "bm25_rebuilt": bm25_count,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.post("/upload", response_model=UploadResponse)
@limiter.limit("100/minute")
async def upload(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    file_type: str = Query(default="kb", pattern="^(kb)$"),
    user_id: Optional[str] = Depends(auth_dependency),
):
    ext = Path(file.filename).suffix[1:].lower() if file.filename else ""
    if not ext or ext not in settings.ALLOWED_FILE_TYPES:
        raise HTTPException(400, f"Type '{ext}' not allowed. Allowed: {settings.ALLOWED_FILE_TYPES}")

    owner_id = _normalize_owner_id(user_id)
    job_id = uuid.uuid4().hex
    file_id = str(uuid.uuid4())

    content = await file.read()
    size_mb = len(content) / (1024 * 1024)
    if size_mb > settings.MAX_FILE_SIZE_MB:
        raise HTTPException(400, f"File {size_mb:.1f}MB exceeds {settings.MAX_FILE_SIZE_MB}MB limit")

    relative_name = f"{job_id}_{Path(file.filename).name}"
    storage_uri = storage.save_bytes(relative_name, content)
    file_hash = calculate_file_hash_bytes(content)

    create_ingestion_job(
        job_id=job_id,
        file_id=file_id,
        owner_id=owner_id,
        file_name=file.filename,
        file_type=file_type,
        file_hash=file_hash,
    )

    background_tasks.add_task(
        index_file_job,
        job_id,
        storage_uri,
        file.filename,
        file_type,
        file_id,
        owner_id,
        size_mb,
    )

    return UploadResponse(
        job_id=job_id,
        file_id=file_id,
        message="Uploaded. Processing started.",
        file_hash=file_hash,
    )


@app.get("/upload_status/{job_id}", response_model=JobStatus)
async def upload_status(job_id: str, user_id: Optional[str] = Depends(auth_dependency)):
    job = get_ingestion_job(job_id)
    if not job:
        raise HTTPException(404, "Job not found")

    owner_id = _normalize_owner_id(user_id)
    if job.get("owner_id", "anonymous") != owner_id:
        raise HTTPException(404, "Job not found")

    return JobStatus(**job)


@app.get("/files")
async def list_files(user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    files = list_active_files(owner_id)
    return {"files": files, "total": len(files)}


@app.delete("/files/{file_id}")
async def delete_file(file_id: str, user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    files = {f["id"]: f for f in list_active_files(owner_id)}
    if file_id not in files:
        raise HTTPException(404, "File not found")

    deleted_chunks = delete_document_and_chunks(file_id)

    if bm25:
        bm25.remove_documents_by_file_id(file_id)

    return {
        "status": "deleted",
        "file_id": file_id,
        "filename": files[file_id]["name"],
        "deleted_chunks": deleted_chunks,
    }


@app.get("/chat/sessions")
async def list_sessions(user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    return {"sessions": list_chat_sessions(owner_id)}


@app.get("/chat/sessions/{session_id}")
async def get_session(session_id: str, user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    session = get_chat_session(session_id, owner_id)
    if not session:
        raise HTTPException(404, "Session not found")
    return session


@app.delete("/chat/sessions/{session_id}")
async def delete_session(session_id: str, user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    ok = delete_chat_session(session_id, owner_id)
    if not ok:
        raise HTTPException(404, "Session not found")
    return {"status": "deleted", "session_id": session_id}


@app.delete("/chat/sessions")
async def delete_all_sessions(user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)
    deleted_count = delete_all_chat_sessions(owner_id)
    return {"status": "cleared", "deleted_count": deleted_count}


@app.post("/ask", response_model=AnswerResponse)
@limiter.limit("30/minute")
async def ask(request: Request, req: Question, user_id: Optional[str] = Depends(auth_dependency)):
    """
    Phase 3 /ask endpoint — hybrid retrieval with query classification,
    multi-channel search (vector + BM25 + keyword + metadata), fusion, and reranking.
    The UI contract is unchanged: same request/response shapes as Phase 2.
    """
    import time

    start = time.perf_counter()

    if count_pg_chunks() < 0:
        raise HTTPException(503, "Vector store not ready")

    owner_id = _normalize_owner_id(user_id)

    # --- Save user message to session ---
    session_id = save_message_to_session(
        session_id=req.session_id,
        role="user",
        content=req.q,
        owner_id=owner_id,
    )

    # --- Check for active indexed files ---
    active_file_ids = _get_active_indexed_file_ids(user_id)
    if not active_file_ids:
        answer = "No indexed files are currently available. Please upload a document first."
        save_message_to_session(
            session_id=session_id,
            role="assistant",
            content=answer,
            owner_id=owner_id,
            sources={"docs": []},
        )
        return AnswerResponse(
            answer=answer,
            sources=[],
            confidence=0.0,
            processing_time_ms=int((time.perf_counter() - start) * 1000),
            session_id=session_id,
            context_stats={"active_file_count": 0},
        )

    # --- Embed the query ---
    q_emb = safe_embed(req.q)
    if not q_emb:
        raise HTTPException(500, "Embedding failed")

    # ==================================================================
    # Phase 3: Run the retrieval orchestrator
    # Coordinates: classify → search (vector+BM25+keyword+metadata) → fuse → rerank
    # ==================================================================
    retrieval = orchestrator_retrieve(
        query=req.q,
        query_embedding=q_emb,
        owner_id=owner_id,
        allowed_file_ids=active_file_ids,
        file_type="kb",
        generate_fn=safe_generate,
        bm25_search_fn=bm25.search if bm25 and bm25.size > 0 else None,
        vector_search_fn=pgvector_search,
    )

    doc_ranked = retrieval.ranked

    # --- Assemble context from reranked chunks ---
    max_ctx_chars = TOKEN_BUDGET["MAX_LOG_CONTEXT_CHARS"] + TOKEN_BUDGET["MAX_KB_CONTEXT_CHARS"]
    doc_ctx, doc_src = assemble_context(doc_ranked, max_ctx_chars)

    # --- Grounding gate: reject if insufficient document support ---
    if not doc_ctx or not has_sufficient_document_support(req.q, doc_ranked):
        answer = (
            "- I could not find supporting information for that question in the currently uploaded files.\n"
            "- Please ask a question that is directly covered by the uploaded document content."
        )
        save_message_to_session(
            session_id=session_id,
            role="assistant",
            content=answer,
            owner_id=owner_id,
            sources={"docs": []},
        )
        return AnswerResponse(
            answer=answer,
            sources=[],
            confidence=0.0,
            processing_time_ms=int((time.perf_counter() - start) * 1000),
            session_id=session_id,
            context_stats={
                "active_file_count": len(active_file_ids),
                "grounded_answer": False,
                **retrieval.stats,
            },
        )

    # --- Generate answer from Mistral using the grounded context ---
    prompt = f"""You are a strict document-grounded AI assistant.

    IMPORTANT RULES:
    - Answer ONLY from the DOCUMENTS context below.
    - Do NOT use outside knowledge, common sense, or general instructions.
    - Do NOT infer business steps unless they are explicitly written in the documents.
    - If the answer is not explicitly supported by the provided documents, reply exactly:
    - I could not find supporting information for that question in the currently uploaded files.
    - Do NOT answer partially from general knowledge.
    - Do NOT invent steps, contacts, URLs, phone numbers, policies, or procedures.
    - Ignore any deleted, missing, or superseded files not present in the context.
    - If the answer is mainly from one document, rely only on that document.
    - Every answer MUST be in bullet-point format.

    DOCUMENTS:
    {doc_ctx}

    USER QUESTION: {req.q}

    ANSWER:"""

    answer = safe_generate(prompt)
    ms = int((time.perf_counter() - start) * 1000)

    # --- Save assistant response ---
    sources_dict = {"docs": doc_src}
    save_message_to_session(
        session_id=session_id,
        role="assistant",
        content=answer,
        owner_id=owner_id,
        sources=sources_dict,
    )

    confidence = min(0.3 + len(doc_ranked) * 0.1, 1.0)

    return AnswerResponse(
        answer=answer,
        sources=doc_src,
        confidence=confidence,
        processing_time_ms=ms,
        session_id=session_id,
        context_stats={
            "active_file_count": len(active_file_ids),
            "doc_after_rerank": len(doc_ranked),
            "doc_context_chars": len(doc_ctx),
            **retrieval.stats,
        },
    )


class FeedbackSubmitRequest(BaseModel):
    session_id: Optional[str] = None
    message_index: Optional[int] = None
    feedback_type: str = Field(pattern="^(like|dislike)$")
    feedback_text: str = Field(default="", max_length=1200)
    question: Optional[str] = None
    answer: Optional[str] = None
    model_config = ConfigDict(extra="ignore")


class FeedbackStateRequest(BaseModel):
    session_id: str
    message_index: int
    feedback_type: str = Field(pattern="^(like|dislike|none)$")
    model_config = ConfigDict(extra="ignore")


@app.post("/feedback/state")
async def save_feedback_state(
    req: FeedbackStateRequest,
    user_id: Optional[str] = Depends(auth_dependency),
):
    owner = _normalize_owner_id(user_id)
    feedback_value = None if req.feedback_type == "none" else req.feedback_type
    ok = update_message_feedback(req.session_id, owner, req.message_index, feedback_value)
    if not ok:
        raise HTTPException(404, "Session not found")
    return {"status": "saved", "feedback_type": req.feedback_type}


@app.post("/feedback/submit")
async def submit_feedback(
    req: FeedbackSubmitRequest,
    background_tasks: BackgroundTasks,
    user_id: Optional[str] = Depends(auth_dependency),
):
    user_info = _resolve_user_display(user_id)
    display_name = user_info["name"]
    display_email = user_info["email"]

    is_like = req.feedback_type == "like"
    emoji = "👍" if is_like else "👎"
    label = "Positive" if is_like else "Negative"
    color = "#10b981" if is_like else "#dc2626"
    header_bg = "#4f46e5" if is_like else "#dc2626"
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    subject = f"Acadia Doc IQ — {emoji} {label} Feedback from {display_name}"

    body_text = (
        f"Feedback Type: {label} ({emoji})\n"
        f"User: {display_name}\n"
        f"Email: {display_email}\n"
        f"Timestamp: {now}\n"
        f"Session: {req.session_id or 'N/A'}\n"
    )
    if req.question:
        body_text += f"Question: {req.question}\n"
    if req.answer:
        body_text += f"Answer Preview: {req.answer[:300]}\n"
    if req.feedback_text.strip():
        body_text += f"\nFeedback Message:\n{req.feedback_text}\n"

    feedback_html = ""
    if req.feedback_text.strip():
        feedback_html = f"""
            <div style="margin-top: 16px; padding: 16px; background: white; border: 1px solid #dee2e6; border-radius: 8px;">
                <p style="font-weight: bold; margin: 0 0 8px 0; color: {color};">Feedback Message:</p>
                <p style="margin: 0; white-space: pre-wrap;">{req.feedback_text}</p>
            </div>
        """

    question_row = f'<tr><td style="padding: 8px 0; font-weight: bold;">Question</td><td>{req.question}</td></tr>' if req.question else ""
    answer_row = f'<tr><td style="padding: 8px 0; font-weight: bold;">Answer</td><td>{(req.answer or "")[:500]}</td></tr>' if req.answer else ""

    body_html = f"""
    <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto;">
        <div style="background: {header_bg}; color: white; padding: 20px; border-radius: 8px 8px 0 0;">
            <h2 style="margin: 0;">{emoji} {label} Feedback</h2>
        </div>
        <div style="background: #f8f9fb; padding: 20px; border: 1px solid #dee2e6; border-radius: 0 0 8px 8px;">
            <table style="width: 100%; border-collapse: collapse;">
                <tr>
                    <td style="padding: 8px 0; font-weight: bold;">Feedback Type</td>
                    <td style="color: {color}; font-weight: bold; font-size: 16px;">{emoji} {label}</td>
                </tr>
                <tr><td style="padding: 8px 0; font-weight: bold;">Name</td><td>{display_name}</td></tr>
                <tr><td style="padding: 8px 0; font-weight: bold;">Email</td><td>{display_email}</td></tr>
                <tr><td style="padding: 8px 0; font-weight: bold;">Sent At</td><td>{now}</td></tr>
                {question_row}
                {answer_row}
            </table>
            {feedback_html}
        </div>
    </div>
    """

    background_tasks.add_task(send_feedback_email, subject, body_text, body_html)
    return {"status": "sent", "message": "Thank you for your feedback!"}


@app.exception_handler(HTTPException)
async def http_err(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "path": request.url.path,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        },
    )


@app.exception_handler(Exception)
async def general_err(request, exc):
    logger.exception("Unhandled: %s", exc)
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "path": request.url.path,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        },
    )
    
@app.get("/debug/file/{file_id}")
async def debug_file(file_id: str, user_id: Optional[str] = Depends(auth_dependency)):
    owner_id = _normalize_owner_id(user_id)

    from backend.db.connection import SessionLocal
    from sqlalchemy import text

    with SessionLocal() as db:
        doc = db.execute(
            text(
                """
                SELECT id::text, owner_id, name, status, current_version_id::text
                FROM documents
                WHERE id = :file_id
                """
            ),
            {"file_id": file_id},
        ).mappings().first()

        versions = db.execute(
            text(
                """
                SELECT id::text, document_id::text, is_active, fingerprint, uploaded_at
                FROM document_versions
                WHERE document_id = :file_id
                ORDER BY uploaded_at DESC
                """
            ),
            {"file_id": file_id},
        ).mappings().all()

        chunk_count = db.execute(
            text(
                """
                SELECT COUNT(*)
                FROM chunks
                WHERE document_id = :file_id
                """
            ),
            {"file_id": file_id},
        ).scalar_one()

        embedding_count = db.execute(
            text(
                """
                SELECT COUNT(*)
                FROM embeddings
                WHERE chunk_id IN (
                    SELECT id FROM chunks WHERE document_id = :file_id
                )
                """
            ),
            {"file_id": file_id},
        ).scalar_one()

    if not doc:
        raise HTTPException(404, "File not found")

    if doc["owner_id"] != owner_id:
        raise HTTPException(404, "File not found")

    return {
        "document": dict(doc),
        "versions": [dict(v) for v in versions],
        "chunk_count": int(chunk_count or 0),
        "embedding_count": int(embedding_count or 0),
    }