"""
Structured Parser — Smart Adaptive Chunking.
Tries 3 detection strategies in order:
  1. Word heading styles (Heading 1/2/3) — fastest, free
  2. Content-based patterns (Scenario A:, Troubleshooting Runbook:) — fast, free
  3. LLM-based section discovery via Haiku — slowest, ~$0.002/doc, works on anything
Only falls back to LLM when strategies 1+2 find zero headings.
"""

from __future__ import annotations

import json
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

import fitz
from docx import Document

from backend.config import settings
from backend.metadata.structure_config import match_operational_section

logger = logging.getLogger("acadia-log-iq")


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------
@dataclass
class ParsedBlock:
    text: str
    block_type: str
    heading: Optional[str] = None
    page_number: Optional[int] = None
    source_order: int = 0
    metadata: Dict = field(default_factory=dict)


@dataclass
class ParsedChunk:
    chunk_index: int
    text: str
    chunk_type: str
    section_heading: Optional[str]
    operational_section: Optional[str]
    page_number: Optional[int]
    source_order: int
    token_estimate: int
    metadata: Dict


# ---------------------------------------------------------------------------
# Regex patterns for classification
# ---------------------------------------------------------------------------
_HEADING_RE = re.compile(r"^\s*(#{1,6}\s+.+|[A-Z][A-Z0-9 /:_\-\(\)]{3,})\s*$")
_BULLET_RE = re.compile(r"^\s*([-*•]|\d+\.)\s+")
_CODE_RE = re.compile(
    r"^\s*(\$|>|kubectl |aws |curl |SELECT |INSERT |UPDATE |DELETE "
    r"|GET |POST |apiVersion:|kind:|FROM |WHERE )"
)
_TABLE_HINT_RE = re.compile(r"\s{2,}|\|")

# --- Content-based heading patterns (Strategy 2) ---
# Detects structural headings from text even when Word styles are 'Normal'.
_CONTENT_HEADING_PATTERNS = [
    # Runbook/document titles
    re.compile(r"^Troubleshooting Runbook\s*:", re.I),
    # Scenario headers: "Scenario A: ...", "Scenario 1: ..."
    re.compile(r"^Scenario\s+[A-Z0-9]+\s*:", re.I),
    # Numbered sections: "1. Overview", "3.2 Configuration"
    re.compile(r"^\d+(?:\.\d+)*\s+[A-Z]", re.I),
    # Chapter/Section/Part headers
    re.compile(r"^(?:Chapter|Section|Part)\s+\d+", re.I),
    # Step-based headers: "Step 1:", "Phase 1:"
    re.compile(r"^(?:Step|Phase|Stage)\s+\d+\s*:", re.I),
    # Common doc section titles (standalone lines)
    re.compile(
        r"^(?:Overview|Introduction|Prerequisites|Procedure|Conclusion|Appendix|"
        r"Escalation Matrix|Escalation Criteria|References|Glossary|"
        r"Executive Summary|Background|Scope|Objectives)\s*:?\s*$", re.I
    ),
    # Section dividers (dashes, equals, asterisks)
    re.compile(r"^[-=*]{5,}$"),
    # Problem / Solution / Workaround style (KB articles)
    re.compile(r"^(?:Problem|Solution|Workaround|Root Cause|Resolution|Impact)\s*:?\s*$", re.I),
]

# --- Sub-section labels (stay WITHIN parent chunk, not split boundaries) ---
_SUB_SECTION_PATTERNS = [
    re.compile(
        r"^(?:Alert Signatures|Severity|Incident Summary|Probable Causes|"
        r"Corrective Actions|Diagnostic Steps|Resolution Steps|"
        r"Validation Steps|Validation / Post.Check|Post-Check|"
        r"Escalation Criteria|Commands / Tools|Commands|"
        r"Expected Results|Root Cause|Workaround|"
        r"Affected Systems|Impact Assessment|"
        r"Pre-Conditions|Post-Conditions|Notes|Warning)\s*:?\s*$", re.I
    ),
]


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def estimate_tokens(text: str) -> int:
    return max(1, len(text) // 4)


def normalize_text(text: str) -> str:
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def _is_content_heading(text: str) -> bool:
    """Check if a line is a structural heading based on its content."""
    stripped = text.strip()
    return any(p.match(stripped) for p in _CONTENT_HEADING_PATTERNS)


def _is_sub_section_label(text: str) -> bool:
    """Check if a line is a sub-section label within a scenario."""
    stripped = text.strip()
    return any(p.match(stripped) for p in _SUB_SECTION_PATTERNS)


def _is_section_divider(text: str) -> bool:
    """Check if a line is a visual section divider."""
    return bool(re.match(r"^[-=*]{5,}$", text.strip()))


def classify_line(line: str) -> str:
    text = line.strip()
    if not text:
        return "blank"
    if _is_section_divider(text):
        return "divider"
    if _is_content_heading(text):
        return "heading"
    if _HEADING_RE.match(text):
        return "heading"
    if _BULLET_RE.match(text):
        return "bullet"
    if settings.ENABLE_CODE_BLOCK_DETECTION and _CODE_RE.match(text):
        return "code"
    if settings.ENABLE_TABLE_PARSING and (_TABLE_HINT_RE.search(text) and len(text.split()) >= 3):
        if "|" in text or "  " in text:
            return "table"
    return "paragraph"


# ---------------------------------------------------------------------------
# LLM-based section discovery (Strategy 3 — fallback)
# ---------------------------------------------------------------------------
def _llm_discover_sections(full_text: str) -> List[str]:
    """
    Use Claude Haiku to identify section boundaries in unstructured text.
    Sends the first ~8000 chars and asks Haiku to return a JSON array
    of section heading strings found in the text.
    Cost: ~$0.002 per document. Only called when strategies 1+2 find nothing.

    Returns a list of heading strings that appear verbatim in the text.
    """
    # Import here to avoid circular imports at module load time
    from backend.services.bedrock_haiku import haiku_client

    # Send a preview — enough to find the document's structural pattern
    preview = full_text[:settings.LLM_CHUNK_FALLBACK_PREVIEW_CHARS]

    system = (
        "You identify section boundaries in documents. "
        "Return strict JSON only. No markdown fences. No commentary."
    )

    prompt = f"""Analyze this document text and identify ALL section headings / titles
that represent logical boundaries between different topics or procedures.

Return a JSON array of the EXACT heading strings as they appear in the text.
Include document titles, scenario names, chapter headers, procedure names, etc.
Do NOT include sub-labels like "Probable Causes:" or "Severity:" — only major sections.

Example output: ["Introduction", "Scenario 1: Network Failure", "Appendix A"]

Document text:
{preview}

Section headings (JSON array only):"""

    try:
        result = haiku_client.invoke_json(
            system=system,
            prompt=prompt,
            max_tokens=1024,
        )

        # The result might be a dict with a key, or a raw list
        if isinstance(result, list):
            headings = [str(h).strip() for h in result if str(h).strip()]
        elif isinstance(result, dict):
            # Try common keys
            for key in ("headings", "sections", "section_headings", "results"):
                if key in result and isinstance(result[key], list):
                    headings = [str(h).strip() for h in result[key] if str(h).strip()]
                    break
            else:
                headings = []
        else:
            headings = []

        # Validate: only keep headings that actually appear in the text
        validated = []
        text_lower = full_text.lower()
        for h in headings:
            if h.lower() in text_lower and len(h) > 3:
                validated.append(h)

        logger.info(
            "LLM section discovery: found %d headings from %d candidates (preview=%d chars)",
            len(validated), len(headings), len(preview),
        )
        return validated

    except Exception as exc:
        logger.warning("LLM section discovery failed: %s", exc)
        return []


def _apply_llm_headings_to_blocks(
    blocks: List[ParsedBlock],
    llm_headings: List[str],
) -> List[ParsedBlock]:
    """
    Post-process blocks: upgrade any paragraph block whose text matches
    an LLM-discovered heading to block_type='heading'.
    """
    heading_set = {h.lower().strip() for h in llm_headings}

    updated = []
    current_heading = None

    for block in blocks:
        text_lower = block.text.strip().lower()

        # Check if this block's text matches an LLM-discovered heading
        if block.block_type == "paragraph" and text_lower in heading_set:
            block = ParsedBlock(
                text=block.text.strip(),
                block_type="heading",
                heading=block.text.strip(),
                page_number=block.page_number,
                source_order=block.source_order,
                metadata=block.metadata,
            )
            current_heading = block.text.strip()
        elif block.block_type == "heading":
            current_heading = block.heading or block.text
        else:
            # Update the heading context for non-heading blocks
            if current_heading and not block.heading:
                block.heading = current_heading

        updated.append(block)

    return updated


# ---------------------------------------------------------------------------
# Block extraction from raw text (for PDFs and plain text)
# ---------------------------------------------------------------------------
def extract_blocks_from_text(
    text: str,
    page_number: Optional[int] = None,
) -> List[ParsedBlock]:
    blocks: List[ParsedBlock] = []
    lines = normalize_text(text).splitlines()
    current_heading: Optional[str] = None
    order = 0
    buffer: List[str] = []
    buffer_type: Optional[str] = None

    def flush():
        nonlocal buffer, buffer_type, order
        if not buffer:
            return
        blocks.append(
            ParsedBlock(
                text="\n".join(buffer).strip(),
                block_type=buffer_type or "paragraph",
                heading=current_heading,
                page_number=page_number,
                source_order=order,
            )
        )
        order += 1
        buffer = []
        buffer_type = None

    for line in lines:
        line_type = classify_line(line)

        if line_type == "blank":
            flush()
            continue

        if line_type == "divider":
            flush()
            continue

        if line_type == "heading":
            flush()
            current_heading = line.strip().lstrip("#").strip()
            blocks.append(
                ParsedBlock(
                    text=current_heading,
                    block_type="heading",
                    heading=current_heading,
                    page_number=page_number,
                    source_order=order,
                )
            )
            order += 1
            continue

        if buffer_type and buffer_type != line_type:
            flush()

        buffer_type = line_type
        buffer.append(line)

    flush()
    return blocks


# ---------------------------------------------------------------------------
# File-type-specific parsers
# ---------------------------------------------------------------------------
def parse_pdf(path: Path) -> List[ParsedBlock]:
    blocks: List[ParsedBlock] = []
    with fitz.open(path) as doc:
        for page_idx, page in enumerate(doc, start=1):
            text = page.get_text("text") or ""
            if text.strip():
                blocks.extend(extract_blocks_from_text(text, page_number=page_idx))
    return blocks


def parse_docx(path: Path) -> List[ParsedBlock]:
    """
    Parse a .docx file into structured blocks.

    Strategy order:
    1. Word heading styles (Heading 1/2/3) — check para.style
    2. Content-based patterns — regex on text content
    3. If both find zero headings → LLM fallback (Haiku discovers sections)
    """
    doc = Document(str(path))
    blocks: List[ParsedBlock] = []
    order = 0
    current_heading = None
    heading_count = 0  # Track how many headings we find

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = (para.style.name or "").lower() if para.style else ""

        # --- Detect block type ---
        # Priority 1: Word heading styles
        if "heading" in style or "title" in style:
            block_type = "heading"
            current_heading = text
            heading_count += 1
        # Priority 2: Content-based heading detection
        elif _is_content_heading(text):
            block_type = "heading"
            current_heading = text
            heading_count += 1
        # Priority 3: Section dividers
        elif _is_section_divider(text):
            continue
        # Priority 4: Sub-section labels (stay within chunk)
        elif _is_sub_section_label(text):
            block_type = "sub_heading"
        # Priority 5: Regular content
        elif text.startswith(("-", "*", "•")):
            block_type = "bullet"
        elif settings.ENABLE_CODE_BLOCK_DETECTION and _CODE_RE.match(text):
            block_type = "code"
        else:
            block_type = "paragraph"

        blocks.append(
            ParsedBlock(
                text=text,
                block_type=block_type,
                heading=current_heading,
                page_number=None,
                source_order=order,
            )
        )
        order += 1

    # Parse tables
    if settings.ENABLE_TABLE_PARSING:
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                blocks.append(
                    ParsedBlock(
                        text="\n".join(rows).strip(),
                        block_type="table",
                        heading=current_heading,
                        page_number=None,
                        source_order=order,
                    )
                )
                order += 1

    # ===================================================================
    # Strategy 3: LLM fallback if no headings were found
    # This handles documents with all-Normal styles AND no recognizable
    # content patterns (vendor manuals, free-form reports, etc.)
    # ===================================================================
    if heading_count == 0 and len(blocks) > 5 and settings.ENABLE_LLM_CHUNK_FALLBACK:
        logger.info(
            "No headings detected (styles or patterns) in %d blocks — "
            "falling back to LLM section discovery",
            len(blocks),
        )
        # Reconstruct full text for LLM analysis
        full_text = "\n".join(b.text for b in blocks if b.text.strip())
        llm_headings = _llm_discover_sections(full_text)

        if llm_headings:
            logger.info("LLM found %d section headings, re-tagging blocks", len(llm_headings))
            blocks = _apply_llm_headings_to_blocks(blocks, llm_headings)
            heading_count = sum(1 for b in blocks if b.block_type == "heading")
            logger.info("After LLM re-tagging: %d heading blocks", heading_count)
        else:
            logger.warning("LLM fallback found no headings — chunking will use char limits only")

    logger.info(
        "parse_docx complete: %d blocks, %d headings (strategy: %s)",
        len(blocks), heading_count,
        "word_styles" if heading_count > 0 else "llm_fallback",
    )

    return blocks


def parse_plain(path: Path) -> List[ParsedBlock]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    blocks = extract_blocks_from_text(text)

    # Check if we found any headings
    heading_count = sum(1 for b in blocks if b.block_type == "heading")

    # LLM fallback for plain text files with no detected headings
    if heading_count == 0 and len(blocks) > 5 and settings.ENABLE_LLM_CHUNK_FALLBACK:
        logger.info("No headings in plain text (%d blocks) — trying LLM discovery", len(blocks))
        full_text = "\n".join(b.text for b in blocks if b.text.strip())
        llm_headings = _llm_discover_sections(full_text)
        if llm_headings:
            blocks = _apply_llm_headings_to_blocks(blocks, llm_headings)

    return blocks


def parse_file(path: Path) -> List[ParsedBlock]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        blocks = parse_pdf(path)
        # LLM fallback for PDFs with no headings
        heading_count = sum(1 for b in blocks if b.block_type == "heading")
        if heading_count == 0 and len(blocks) > 5 and settings.ENABLE_LLM_CHUNK_FALLBACK:
            logger.info("No headings in PDF (%d blocks) — trying LLM discovery", len(blocks))
            full_text = "\n".join(b.text for b in blocks if b.text.strip())
            llm_headings = _llm_discover_sections(full_text)
            if llm_headings:
                blocks = _apply_llm_headings_to_blocks(blocks, llm_headings)
        return blocks
    if suffix == ".docx":
        return parse_docx(path)
    return parse_plain(path)


# ---------------------------------------------------------------------------
# Chunk type classification
# ---------------------------------------------------------------------------
def choose_chunk_type(block_types: List[str], heading: Optional[str]) -> str:
    """Classify a chunk's operational type from its blocks and heading."""
    if heading:
        matched = match_operational_section(heading)
        if matched:
            return matched.chunk_type

    heading_lower = (heading or "").lower()
    if "scenario" in heading_lower:
        if any(kw in heading_lower for kw in [
            "troubleshoot", "diagnostic", "failure", "issue", "down",
            "error", "fault", "loss", "degrad",
        ]):
            return "diagnostic_chunk"
        return "general_chunk"
    if "escalation" in heading_lower:
        return "escalation_chunk"

    if "code" in block_types:
        return "command_chunk"
    if "table" in block_types:
        return "validation_chunk"
    return "general_chunk"


# ---------------------------------------------------------------------------
# Scenario-aware chunk builder
# ---------------------------------------------------------------------------
def build_chunks(blocks: List[ParsedBlock]) -> List[ParsedChunk]:
    """
    Build chunks from parsed blocks with scenario-aware boundaries.

    Works with all 3 heading detection strategies:
    - Word styles → headings already tagged
    - Content patterns → headings already tagged
    - LLM fallback → headings re-tagged by _apply_llm_headings_to_blocks

    Each heading starts a new chunk. Sub-section labels stay within their
    parent chunk. Only splits mid-section if content exceeds 6000 chars.
    """
    chunks: List[ParsedChunk] = []
    current_group: List[ParsedBlock] = []
    current_heading: Optional[str] = None
    current_operational_section: Optional[str] = None
    char_count = 0
    chunk_index = 0

    def flush():
        nonlocal current_group, char_count, chunk_index
        if not current_group:
            return

        text_parts = []
        for block in current_group:
            if block.block_type == "heading":
                continue
            text_parts.append(block.text)

        text = "\n\n".join(text_parts).strip()
        if not text:
            current_group = []
            char_count = 0
            return

        # Prepend heading to chunk text for better retrieval
        if current_heading:
            text = f"[{current_heading}]\n\n{text}"

        block_types = [b.block_type for b in current_group]
        chunks.append(
            ParsedChunk(
                chunk_index=chunk_index,
                text=text,
                chunk_type=choose_chunk_type(block_types, current_heading),
                section_heading=current_heading,
                operational_section=current_operational_section,
                page_number=current_group[0].page_number,
                source_order=current_group[0].source_order,
                token_estimate=estimate_tokens(text),
                metadata={
                    "block_types": block_types,
                    "page_numbers": [
                        b.page_number for b in current_group if b.page_number is not None
                    ],
                },
            )
        )
        chunk_index += 1
        current_group = []
        char_count = 0

    for block in blocks:
        # Heading blocks: start a new chunk
        if block.block_type == "heading":
            flush()
            current_heading = block.heading or block.text
            match = match_operational_section(current_heading or "")
            current_operational_section = match.canonical_name if match else None
            current_group.append(block)
            continue

        # Sub-heading blocks: stay within current chunk
        if block.block_type == "sub_heading":
            current_group.append(block)
            char_count += len(block.text) + 2
            continue

        # Regular content blocks
        projected = char_count + len(block.text)
        max_chunk_chars = max(settings.CHUNK_MAX_CHARS, 6000)

        if current_group and projected > max_chunk_chars and char_count >= settings.CHUNK_MIN_CHARS:
            flush()
            if current_heading:
                current_group.append(
                    ParsedBlock(
                        text=current_heading,
                        block_type="heading",
                        heading=current_heading,
                        page_number=block.page_number,
                        source_order=block.source_order,
                    )
                )

        current_group.append(block)
        char_count += len(block.text) + 2

    flush()
    return chunks 