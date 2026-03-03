import json
import logging
import hashlib
import os
import sys
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional, Tuple, TypedDict

import boto3
from botocore.config import Config as BotoConfig
from fastapi import (
    BackgroundTasks, Depends, FastAPI, File, Header,
    HTTPException, Query, Request, UploadFile, status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from slowapi.util import get_remote_address
from pydantic import BaseModel, ConfigDict, Field, field_validator

from config import settings
from vector_store import get_collection, get_bm25_index, rebuild_bm25_from_chroma, reset_chroma_collection


if sys.version_info < (3, 11):
    raise RuntimeError("Python 3.11+ required")


# ============================================================================
# LOGGING
# ============================================================================
logging.basicConfig(
    level=getattr(logging, settings.LOG_LEVEL.upper(), logging.INFO),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
)
logger = logging.getLogger("acadia-log-iq")


# ============================================================================
# TOKEN BUDGET — MISTRAL 7B (32K CONTEXT)
# ============================================================================
# Conservative: 1 token ~ 3 chars for Mistral tokenizer
#
# ┌──────────────────────────────────────────────────────────┐
# │  MISTRAL 7B TOKEN BUDGET (32,768 tokens)                │
# ├──────────────────────────────────────────────────────────┤
# │  System prompt + question       ~500 tokens  (1,500 ch) │
# │  Log context (max)            ~8,000 tokens (24,000 ch) │
# │  KB context (max)             ~8,000 tokens (24,000 ch) │
# │  Generation output (reserved) ~2,048 tokens             │
# │  Safety buffer               ~14,220 tokens             │
# └──────────────────────────────────────────────────────────┘
# ============================================================================
CHARS_PER_TOKEN = 3

TOKEN_BUDGET = {
    "MODEL_MAX_TOKENS": 32_768,
    "MAX_LOG_CONTEXT_CHARS": 24_000,
    "MAX_KB_CONTEXT_CHARS": 24_000,
    "MAX_GENERATION_TOKENS": 2_048,
    "MAX_SINGLE_CHUNK_CHARS": 6_000,
    "PROMPT_OVERHEAD_CHARS": 1_500,
}

_MAX_PROMPT_TOKENS = TOKEN_BUDGET["MODEL_MAX_TOKENS"] - TOKEN_BUDGET["MAX_GENERATION_TOKENS"]
MAX_TOTAL_PROMPT_CHARS = (_MAX_PROMPT_TOKENS * CHARS_PER_TOKEN) - TOKEN_BUDGET["PROMPT_OVERHEAD_CHARS"]

CHUNK_CONFIG = {
    "LINES_PER_CHUNK": 80,
    "MAX_CHUNK_CHARS": 6_000,
    "OVERLAP_LINES": 10,
}

# Hybrid search config
HYBRID_CONFIG = {
    "VECTOR_WEIGHT": 0.6,
    "BM25_WEIGHT": 0.4,
    "VECTOR_CANDIDATES": 15,
    "BM25_CANDIDATES": 15,
    "RERANK_TOP_K_LOG": 6,
    "RERANK_TOP_K_KB": 5,
}


# ============================================================================
# TYPES
# ============================================================================
class JobInfo(TypedDict, total=False):
    job_id: str
    status: str
    processed_chunks: int
    total_chunks: int
    successful_chunks: int
    file: Optional[str]
    file_type: Optional[str]
    file_size_mb: float
    file_hash: str
    error: Optional[str]
    created_at: datetime
    completed_at: Optional[datetime]


UPLOAD_JOBS: Dict[str, JobInfo] = {}
coll = None
bm25 = None


# ============================================================================
# BEDROCK CLIENT
# ============================================================================
def _make_bedrock_client():
    boto_cfg = BotoConfig(
        retries={"max_attempts": 10, "mode": "adaptive"},
        read_timeout=120, connect_timeout=30, tcp_keepalive=True,
    )
    kwargs = {
        "service_name": "bedrock-runtime",
        "region_name": settings.AWS_REGION,
        "config": boto_cfg,
    }
    if settings.AWS_ACCESS_KEY_ID and settings.AWS_SECRET_ACCESS_KEY:
        kwargs["aws_access_key_id"] = settings.AWS_ACCESS_KEY_ID
        kwargs["aws_secret_access_key"] = settings.AWS_SECRET_ACCESS_KEY
        if settings.AWS_SESSION_TOKEN:
            kwargs["aws_session_token"] = settings.AWS_SESSION_TOKEN
    return boto3.client(**kwargs)


bedrock = _make_bedrock_client()


# ============================================================================
# LIFESPAN
# ============================================================================
@asynccontextmanager
async def lifespan(app: FastAPI):
    global coll, bm25
    logger.info("Starting API...")
    coll = get_collection()
    logger.info("Chroma ready: %s", settings.COLLECTION_NAME)
    bm25 = get_bm25_index()
    doc_count = rebuild_bm25_from_chroma(coll)
    logger.info("BM25 ready: %d docs", doc_count)
    yield
    logger.info("Shutting down...")


app = FastAPI(
    title="Acadia's Log IQ API",
    description="AI log analysis — Hybrid Search + Re-ranking",
    version="2.0.0",
    lifespan=lifespan,
)

limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8501", "http://localhost:3000"],
    allow_credentials=True, allow_methods=["*"], allow_headers=["*"],
    expose_headers=["X-Processing-Time"], max_age=600,
)


def verify_api_key(x_api_key: Optional[str] = Header(default=None, alias="X-API-Key")) -> bool:
    if settings.API_KEY:
        if not x_api_key or x_api_key != settings.API_KEY:
            raise HTTPException(status_code=status.HTTP_401_UNAUTHORIZED, detail="Invalid API key")
    return True


# ============================================================================
# PYDANTIC MODELS
# ============================================================================
class Question(BaseModel):
    q: str = Field(min_length=1, max_length=1000)
    model_config = ConfigDict(extra="ignore")
    @field_validator("q")
    @classmethod
    def validate_q(cls, v):
        v = (v or "").strip()
        if not v: raise ValueError("Empty")
        return v

class UploadResponse(BaseModel):
    job_id: str
    message: str
    file_hash: Optional[str] = None
    model_config = ConfigDict(extra="ignore")

class JobStatus(BaseModel):
    job_id: str
    status: str
    processed_chunks: int = 0
    total_chunks: Optional[int] = None
    successful_chunks: Optional[int] = None
    file: Optional[str] = None
    file_type: Optional[str] = None
    file_size_mb: Optional[float] = None
    file_hash: Optional[str] = None
    error: Optional[str] = None
    created_at: datetime
    completed_at: Optional[datetime] = None
    model_config = ConfigDict(extra="ignore")

class AnswerResponse(BaseModel):
    answer: str
    log_sources: List[str]
    kb_sources: List[str]
    confidence: float = Field(ge=0, le=1)
    processing_time_ms: Optional[int] = None
    context_stats: Optional[Dict] = None
    model_config = ConfigDict(extra="ignore")


# ============================================================================
# EMBEDDING
# ============================================================================
def safe_embed(text: str) -> Optional[List[float]]:
    if not text or not text.strip():
        return None
    try:
        body = json.dumps({"inputText": text[: settings.MAX_CHARS]}).encode("utf-8")
        resp = bedrock.invoke_model(
            modelId=settings.BEDROCK_EMBED_MODEL, body=body,
            accept="application/json", contentType="application/json",
        )
        payload = json.loads(resp["body"].read().decode("utf-8"))
        emb = payload.get("embedding")
        return emb if isinstance(emb, list) else None
    except Exception as e:
        logger.exception("Embed failed: %s", e)
        return None


# ============================================================================
# LLM GENERATION — MISTRAL 7B (32K) WITH HARD TRUNCATION
# ============================================================================
def estimate_tokens(text: str) -> int:
    return len(text) // CHARS_PER_TOKEN


def safe_generate(prompt: str, max_tokens: int = None) -> str:
    """Mistral 7B generation with Layer 3 hard truncation safety."""
    if max_tokens is None:
        max_tokens = TOKEN_BUDGET["MAX_GENERATION_TOKENS"]
    try:
        max_prompt_tokens = TOKEN_BUDGET["MODEL_MAX_TOKENS"] - max_tokens - 200
        max_prompt_chars = max_prompt_tokens * CHARS_PER_TOKEN

        if len(prompt) > max_prompt_chars:
            logger.warning("TRUNCATING: %d -> %d chars", len(prompt), max_prompt_chars)
            marker = "\nANSWER:"
            pos = prompt.rfind(marker)
            if pos > 0:
                tail = prompt[pos:]
                prompt = prompt[: max_prompt_chars - len(tail) - 80] + \
                    "\n\n[... context truncated ...]\n" + tail
            else:
                prompt = prompt[:max_prompt_chars] + "\n\n[... truncated ...]\n"

        logger.info("safe_generate: %d chars (~%d tok)", len(prompt), estimate_tokens(prompt))

        body = json.dumps({
            "prompt": prompt,
            "max_tokens": max_tokens,
            "temperature": 0.1,
            "top_p": 0.9,
        }).encode("utf-8")

        resp = bedrock.invoke_model(
            modelId=settings.BEDROCK_LLM_MODEL, body=body,
            accept="application/json", contentType="application/json",
        )
        payload = json.loads(resp["body"].read().decode("utf-8"))

        if isinstance(payload, dict):
            if "outputs" in payload and payload["outputs"]:
                return (payload["outputs"][0].get("text") or "").strip() or "No response."
            if "generation" in payload:
                return str(payload["generation"]).strip()
            if "outputText" in payload:
                return str(payload["outputText"]).strip()
        return "No response generated."
    except Exception as e:
        logger.exception("Generation failed: %s", e)
        return "Error generating response. Please try again."


# ============================================================================
# HYBRID SEARCH: VECTOR + BM25 + RECIPROCAL RANK FUSION
# ============================================================================
def hybrid_search(
    query: str, query_embedding: List[float], file_type: str, n_results: int = 10,
) -> List[Tuple[str, str, Dict, float]]:
    """
    Merge ChromaDB vector results and BM25 keyword results using
    Reciprocal Rank Fusion (RRF).
    """
    v_weight = HYBRID_CONFIG["VECTOR_WEIGHT"]
    b_weight = HYBRID_CONFIG["BM25_WEIGHT"]
    rrf_k = 60

    # Vector search
    vector_results = {}
    try:
        cr = coll.query(
            query_embeddings=[query_embedding],
            n_results=HYBRID_CONFIG["VECTOR_CANDIDATES"],
            where={"file_type": file_type},
        )
        for i, doc_id in enumerate((cr.get("ids") or [[]])[0]):
            docs = (cr.get("documents") or [[]])[0]
            metas = (cr.get("metadatas") or [[]])[0]
            dists = (cr.get("distances") or [[]])[0]
            vector_results[doc_id] = {
                "text": docs[i] if i < len(docs) else "",
                "metadata": metas[i] if i < len(metas) else {},
                "rank": i + 1,
                "similarity": max(0, 1 - (dists[i] if i < len(dists) else 1.0)),
            }
    except Exception as e:
        logger.warning("Vector search failed (%s): %s", file_type, e)

    # BM25 search
    bm25_results = {}
    try:
        if bm25 and bm25.size > 0:
            for i, (doc_id, text, meta, score) in enumerate(
                bm25.search(query, n_results=HYBRID_CONFIG["BM25_CANDIDATES"], file_type=file_type)
            ):
                bm25_results[doc_id] = {
                    "text": text, "metadata": meta,
                    "rank": i + 1, "bm25_score": score,
                }
    except Exception as e:
        logger.warning("BM25 search failed: %s", e)

    # RRF merge
    combined: Dict[str, float] = {}
    all_data: Dict[str, dict] = {}

    for doc_id, data in vector_results.items():
        combined[doc_id] = combined.get(doc_id, 0) + v_weight / (rrf_k + data["rank"])
        all_data[doc_id] = data

    for doc_id, data in bm25_results.items():
        combined[doc_id] = combined.get(doc_id, 0) + b_weight / (rrf_k + data["rank"])
        if doc_id not in all_data:
            all_data[doc_id] = data

    sorted_ids = sorted(combined, key=lambda x: combined[x], reverse=True)

    results = [
        (did, all_data[did]["text"], all_data[did]["metadata"], combined[did])
        for did in sorted_ids[:n_results]
    ]

    logger.info(
        "Hybrid(%s): vec=%d, bm25=%d, merged=%d, returned=%d",
        file_type, len(vector_results), len(bm25_results), len(combined), len(results),
    )
    return results


# ============================================================================
# RE-RANKING VIA LLM
# ============================================================================
def rerank_chunks(
    query: str, chunks: List[Tuple[str, str, Dict, float]], top_k: int = 6,
) -> List[Tuple[str, str, Dict, float]]:
    """
    Re-rank retrieved chunks by asking Mistral 7B to score each for relevance.
    Single LLM call scores all candidates. Falls back to hybrid scores on failure.
    """
    if not chunks or len(chunks) <= 1:
        return chunks[:top_k]

    # Limit candidates to avoid blowing the re-rank prompt budget
    candidates = chunks[:min(len(chunks), 12)]

    previews = []
    for i, (_, text, meta, _) in enumerate(candidates):
        preview = text[:600].replace("\n", " ").strip()
        src = meta.get("source", "?")
        previews.append(f"[{i+1}] (source: {src}) {preview}")

    rerank_prompt = f"""Rate each chunk's relevance to the question (0=irrelevant, 10=perfect match).
Question: {query}

Chunks:
{chr(10).join(previews)}

Respond ONLY with a JSON array: [{{"chunk":1,"score":8}}, ...]"""

    try:
        resp = safe_generate(rerank_prompt, max_tokens=512)

        # Extract JSON
        raw = resp.strip()
        if "```" in raw:
            raw = raw.split("```")[1] if "```json" not in raw else raw.split("```json")[1].split("```")[0]
        start, end = raw.find("["), raw.rfind("]") + 1
        if start < 0 or end <= start:
            raise ValueError("No JSON array found")

        scores = json.loads(raw[start:end])

        scored = []
        for item in scores:
            idx = item.get("chunk", 0) - 1
            relevance = item.get("score", 0)
            if 0 <= idx < len(candidates):
                did, text, meta, orig = candidates[idx]
                final = (relevance / 10.0) * 0.7 + orig * 30 * 0.3
                scored.append((did, text, meta, final))

        # Add unscored candidates
        scored_ids = {s[0] for s in scored}
        for c in candidates:
            if c[0] not in scored_ids:
                scored.append(c)

        scored.sort(key=lambda x: x[3], reverse=True)
        logger.info("Re-ranked %d -> top %d", len(candidates), top_k)
        return scored[:top_k]

    except Exception as e:
        logger.warning("Re-rank failed (%s), using hybrid scores", e)
        return candidates[:top_k]


# ============================================================================
# CONTEXT ASSEMBLY
# ============================================================================
def truncate_chunk(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    t = text[:max_chars]
    nl = t.rfind("\n")
    if nl > max_chars * 0.7:
        t = t[:nl]
    return t + "\n[... truncated ...]"


def assemble_context(
    ranked: List[Tuple[str, str, Dict, float]], max_total_chars: int,
) -> Tuple[str, List[str]]:
    """Build context from ranked chunks. Returns (context_str, sources)."""
    max_chunk = TOKEN_BUDGET["MAX_SINGLE_CHUNK_CHARS"]
    if not ranked:
        return "", []

    parts, sources, total = [], set(), 0
    for i, (_, text, meta, score) in enumerate(ranked):
        if not text or not text.strip():
            continue
        chunk = truncate_chunk(text.strip(), max_chunk)
        src = meta.get("source", "unknown")
        entry = f"[Source: {src}]\n{chunk}"
        if total + len(entry) > max_total_chars:
            remaining = max_total_chars - total
            if remaining > 300:
                parts.append(f"[Source: {src}]\n{truncate_chunk(chunk, remaining - 60)}")
                sources.add(src)
            break
        parts.append(entry)
        sources.add(src)
        total += len(entry)

    return "\n\n".join(parts), sorted(s for s in sources if s)


# ============================================================================
# FILE HASHING
# ============================================================================
def calculate_file_hash(fp: Path) -> str:
    sha = hashlib.sha256()
    with fp.open("rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            sha.update(chunk)
    return sha.hexdigest()


# ============================================================================
# TEXT EXTRACTION (PDF, DOCX)
# ============================================================================
def extract_text_from_pdf(fp: Path) -> str:
    try:
        import fitz
        parts = []
        with fitz.open(str(fp)) as doc:
            for i, page in enumerate(doc):
                t = page.get_text("text")
                if t and t.strip():
                    parts.append(f"--- Page {i+1} ---\n{t}")
        if parts:
            logger.info("PDF: %d pages, %d chars from %s", len(parts), sum(len(p) for p in parts), fp.name)
            return "\n\n".join(parts)
        return ""
    except ImportError:
        pass
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(fp)) as pdf:
            for i, page in enumerate(pdf.pages):
                t = page.extract_text()
                if t and t.strip():
                    parts.append(f"--- Page {i+1} ---\n{t}")
        return "\n\n".join(parts) if parts else ""
    except ImportError:
        raise RuntimeError("No PDF library. pip install PyMuPDF")


def extract_text_from_docx(fp: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        raise RuntimeError("pip install python-docx")
    doc = Document(str(fp))
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for ti, table in enumerate(doc.tables):
        rows = [" | ".join(c.text.strip() for c in r.cells) for r in table.rows]
        if rows:
            parts.append(f"--- Table {ti+1} ---")
            parts.extend(rows)
    return "\n".join(parts)


def extract_text(fp: Path) -> str:
    ext = fp.suffix.lower()
    if ext == ".pdf": return extract_text_from_pdf(fp)
    if ext == ".docx": return extract_text_from_docx(fp)
    return fp.read_text(encoding="utf-8", errors="ignore")


# ============================================================================
# CHUNKING
# ============================================================================
def iter_text_chunks(text, max_chars=None, lines_per=None, overlap=None):
    max_chars = max_chars or CHUNK_CONFIG["MAX_CHUNK_CHARS"]
    lines_per = lines_per or CHUNK_CONFIG["LINES_PER_CHUNK"]
    overlap = overlap or CHUNK_CONFIG["OVERLAP_LINES"]
    lines = text.splitlines(keepends=True)
    if not lines: return
    buf, sz = [], 0
    for ln in lines:
        buf.append(ln); sz += len(ln)
        if len(buf) >= lines_per or sz >= max_chars:
            yield "".join(buf)
            if overlap > 0 and len(buf) > overlap:
                buf = buf[-overlap:]
                sz = sum(len(l) for l in buf)
            else:
                buf.clear(); sz = 0
    if buf:
        yield "".join(buf)


def iter_line_chunks(fp: Path, lines_per=None):
    lines_per = lines_per or CHUNK_CONFIG["LINES_PER_CHUNK"]
    ext = fp.suffix.lower()
    if ext in (".pdf", ".docx"):
        try:
            text = extract_text(fp)
            if not text or not text.strip(): return
            logger.info("Extracted %d chars from %s", len(text), fp.name)
            n = 0
            for c in iter_text_chunks(text, lines_per=lines_per):
                n += 1; yield c
            logger.info("%d chunks from %s", n, fp.name)
        except RuntimeError as e:
            logger.error("Extraction: %s", e); raise
        except Exception as e:
            logger.exception("Failed: %s", e)
        return
    try:
        text = fp.read_text(encoding="utf-8", errors="ignore")
        if text and text.strip():
            for c in iter_text_chunks(text, lines_per=lines_per):
                yield c
    except Exception as e:
        logger.exception("Read failed %s: %s", fp, e)


# ============================================================================
# INDEX JOB — UPDATES BOTH CHROMA + BM25
# ============================================================================
async def index_file_job(job_id: str, fp: Path, filename: str, file_type: str):
    try:
        job = UPLOAD_JOBS.get(job_id)
        if not job: return
        job["status"] = "running"
        total, ok = 0, 0
        b_emb, b_doc, b_meta, b_ids = [], [], [], []
        bm25_ids, bm25_docs, bm25_metas = [], [], []

        for chunk in iter_line_chunks(fp):
            if not chunk or not chunk.strip(): continue
            total += 1
            emb = safe_embed(chunk)
            if not emb: continue
            cid = f"{job_id}-{ok}"
            ok += 1
            meta = {
                "source": filename, "file_type": file_type,
                "job_id": job_id, "chunk_index": ok - 1,
                "char_count": len(chunk),
                "timestamp": datetime.now(timezone.utc).isoformat(),
            }
            b_ids.append(cid); b_emb.append(emb)
            b_doc.append(chunk); b_meta.append(meta)
            bm25_ids.append(cid); bm25_docs.append(chunk); bm25_metas.append(meta)

            if ok % 10 == 0: job["processed_chunks"] = ok
            if len(b_emb) >= settings.BATCH_SIZE:
                coll.add(ids=b_ids, embeddings=b_emb, metadatas=b_meta, documents=b_doc)
                b_ids, b_emb, b_meta, b_doc = [], [], [], []

        if b_emb:
            coll.add(ids=b_ids, embeddings=b_emb, metadatas=b_meta, documents=b_doc)

        if bm25 and bm25_ids:
            bm25.add_documents_batch(bm25_ids, bm25_docs, bm25_metas)
            logger.info("BM25 +%d (total: %d)", len(bm25_ids), bm25.size)

        try: fp.unlink(missing_ok=True)
        except: pass

        job.update({"status": "done", "processed_chunks": ok,
                     "total_chunks": total, "successful_chunks": ok,
                     "completed_at": datetime.now(timezone.utc)})
        logger.info("Job %s DONE: %d/%d from %s", job_id, ok, total, filename)

    except Exception as e:
        logger.exception("Job failed %s: %s", job_id, e)
        if job_id in UPLOAD_JOBS:
            UPLOAD_JOBS[job_id].update({
                "status": "failed", "error": str(e),
                "completed_at": datetime.now(timezone.utc),
            })
        try: fp.unlink(missing_ok=True)
        except: pass


# ============================================================================
# MIDDLEWARE
# ============================================================================
@app.middleware("http")
async def log_requests(request: Request, call_next):
    import time
    start = time.perf_counter()
    response = await call_next(request)
    ms = (time.perf_counter() - start) * 1000.0
    response.headers["X-Processing-Time"] = f"{ms:.2f}ms"
    logger.info("%s %s -> %s (%.2fms)", request.method, request.url.path, response.status_code, ms)
    return response


# ============================================================================
# ROUTES
# ============================================================================
@app.get("/health")
async def health_check():
    cc = 0
    try:
        if coll: cc = coll.count()
    except: pass
    return {
        "status": "healthy",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "model": settings.BEDROCK_LLM_MODEL,
        "services": {
            "vector_store": f"{cc} chunks" if coll else "uninitialized",
            "bm25_index": f"{bm25.size} docs" if bm25 else "uninitialized",
            "bedrock": "available",
        },
        "search_mode": "hybrid (vector + BM25 + re-ranking)",
        "token_budget": {
            "model_max": TOKEN_BUDGET["MODEL_MAX_TOKENS"],
            "max_log_chars": TOKEN_BUDGET["MAX_LOG_CONTEXT_CHARS"],
            "max_kb_chars": TOKEN_BUDGET["MAX_KB_CONTEXT_CHARS"],
        },
    }


@app.post("/reset")
async def reset_all(_: bool = Depends(verify_api_key)):
    """
    FULL RESET — safely clears all data without filesystem corruption.

    Uses ChromaDB's own delete_collection API (NOT shutil.rmtree) to avoid:
    - "Device or resource busy" (can't rmtree a Docker mount point)
    - "readonly database" (corrupted SQLite WAL after partial deletion)

    Steps:
    1. ChromaDB: delete_collection + recreate (same SQLite connection)
    2. BM25: clear in-memory index
    3. Uploads: delete files from EC2 volume
    4. Jobs: clear in-memory tracking
    """
    global coll, bm25
    errors = []

    # 1. Reset ChromaDB via API (safe — no filesystem deletion)
    deleted_chunks = 0
    try:
        deleted_chunks = reset_chroma_collection()
        # Get fresh collection reference
        coll = get_collection()
        logger.info("RESET: ChromaDB cleared (%d chunks deleted), fresh collection ready", deleted_chunks)
    except Exception as e:
        logger.exception("RESET: ChromaDB reset failed: %s", e)
        errors.append(f"ChromaDB: {e}")
        # Try to at least get a working collection reference
        try:
            coll = get_collection()
        except Exception:
            pass

    # 2. Clear BM25 keyword index
    try:
        if bm25:
            old_size = bm25.size
            bm25.clear()
            logger.info("RESET: BM25 cleared (%d docs)", old_size)
    except Exception as e:
        logger.exception("RESET: BM25 clear failed: %s", e)
        errors.append(f"BM25: {e}")

    # 3. Delete uploaded files from EC2 volume
    deleted_files = 0
    try:
        upload_dir = settings.UPLOAD_DIR
        if upload_dir.exists():
            for f in upload_dir.iterdir():
                if f.is_file():
                    try:
                        f.unlink()
                        deleted_files += 1
                    except Exception as fe:
                        logger.warning("RESET: Could not delete %s: %s", f.name, fe)
            logger.info("RESET: Deleted %d files from %s", deleted_files, upload_dir)
    except Exception as e:
        logger.exception("RESET: Upload cleanup failed: %s", e)
        errors.append(f"Uploads: {e}")

    # 4. Clear job tracking
    job_count = len(UPLOAD_JOBS)
    UPLOAD_JOBS.clear()
    logger.info("RESET: Cleared %d job records", job_count)

    # Verify ChromaDB is truly empty
    verify_count = 0
    try:
        if coll:
            verify_count = coll.count()
    except Exception:
        pass

    if errors:
        logger.warning("RESET partial (%d errors): %s", len(errors), errors)
        return {
            "status": "partial",
            "message": f"Reset completed with {len(errors)} error(s)",
            "errors": errors,
            "deleted_chunks": deleted_chunks,
            "deleted_files": deleted_files,
            "cleared_jobs": job_count,
            "chroma_count_after": verify_count,
            "bm25_count_after": bm25.size if bm25 else 0,
            "timestamp": datetime.now(timezone.utc).isoformat(),
        }

    logger.info("RESET: Complete — %d chunks, %d files, %d jobs cleared",
                deleted_chunks, deleted_files, job_count)
    return {
        "status": "success",
        "message": "All data deleted — ChromaDB, BM25, uploads, jobs cleared",
        "deleted_chunks": deleted_chunks,
        "deleted_files": deleted_files,
        "cleared_jobs": job_count,
        "chroma_count_after": verify_count,
        "bm25_count_after": 0,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    }


@app.post("/upload", response_model=UploadResponse)
@limiter.limit("10/minute")
async def upload(
    request: Request, background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    file_type: str = Query(default="log", pattern="^(log|kb)$"),
    _: bool = Depends(verify_api_key),
):
    ext = Path(file.filename).suffix[1:].lower() if file.filename else ""
    if not ext or ext not in settings.ALLOWED_FILE_TYPES:
        raise HTTPException(400, f"Type '{ext}' not allowed. Allowed: {settings.ALLOWED_FILE_TYPES}")

    job_id = uuid.uuid4().hex
    fp = settings.UPLOAD_DIR / f"{job_id}_{Path(file.filename).name}"
    content = await file.read()
    size_mb = len(content) / (1024 * 1024)
    if size_mb > settings.MAX_FILE_SIZE_MB:
        raise HTTPException(400, f"File {size_mb:.1f}MB exceeds {settings.MAX_FILE_SIZE_MB}MB limit")

    fp.write_bytes(content)
    fh = calculate_file_hash(fp)

    UPLOAD_JOBS[job_id] = {
        "job_id": job_id, "status": "queued",
        "processed_chunks": 0, "total_chunks": 0, "successful_chunks": 0,
        "file": file.filename, "file_type": file_type,
        "file_size_mb": size_mb, "file_hash": fh,
        "created_at": datetime.now(timezone.utc),
    }
    background_tasks.add_task(index_file_job, job_id, fp, file.filename, file_type)
    logger.info("Upload: %s (%.1fMB, %s)", file.filename, size_mb, file_type)
    return UploadResponse(job_id=job_id, message="Uploaded. Processing started.", file_hash=fh)


@app.get("/upload_status/{job_id}", response_model=JobStatus)
async def upload_status(job_id: str, _: bool = Depends(verify_api_key)):
    job = UPLOAD_JOBS.get(job_id)
    if not job: raise HTTPException(404, "Job not found")
    return JobStatus(**job)


@app.post("/ask", response_model=AnswerResponse)
@limiter.limit("30/minute")
async def ask(request: Request, req: Question, _: bool = Depends(verify_api_key)):
    """
    RAG pipeline: Hybrid Search → Re-Rank → Budget Assembly → Generate.

    1. Embed question
    2. Hybrid search: Vector (ChromaDB) + Keyword (BM25) merged via RRF
    3. Re-rank: LLM scores each chunk for relevance
    4. Assemble context within 32K token budget
    5. Generate answer with Mistral 7B
    """
    import time
    start = time.perf_counter()

    if not coll:
        raise HTTPException(503, "Vector store not ready")

    q_emb = safe_embed(req.q)
    if not q_emb:
        raise HTTPException(500, "Embedding failed")

    # 2. HYBRID SEARCH
    log_cands = hybrid_search(req.q, q_emb, "log", HYBRID_CONFIG["VECTOR_CANDIDATES"])
    kb_cands = hybrid_search(req.q, q_emb, "kb", HYBRID_CONFIG["BM25_CANDIDATES"])

    # 3. RE-RANK
    log_ranked = rerank_chunks(req.q, log_cands, HYBRID_CONFIG["RERANK_TOP_K_LOG"])
    kb_ranked = rerank_chunks(req.q, kb_cands, HYBRID_CONFIG["RERANK_TOP_K_KB"])

    # 4. ASSEMBLE (Layer 1)
    log_ctx, log_src = assemble_context(log_ranked, TOKEN_BUDGET["MAX_LOG_CONTEXT_CHARS"])
    kb_ctx, kb_src = assemble_context(kb_ranked, TOKEN_BUDGET["MAX_KB_CONTEXT_CHARS"])

    if not log_ctx: log_ctx = "No relevant log entries found."
    if not kb_ctx: kb_ctx = "No relevant knowledge base articles found."

    # Layer 2: combined check
    combined = len(log_ctx) + len(kb_ctx)
    max_comb = MAX_TOTAL_PROMPT_CHARS - TOKEN_BUDGET["PROMPT_OVERHEAD_CHARS"]
    if combined > max_comb:
        half = max_comb // 2
        if len(log_ctx) > half: log_ctx = truncate_chunk(log_ctx, half)
        if len(kb_ctx) > half: kb_ctx = truncate_chunk(kb_ctx, half)

    confidence = min(0.3 + len(log_ranked) * 0.08 + len(kb_ranked) * 0.1, 1.0)

    prompt = f"""You are an expert system administrator analyzing logs and knowledge base articles.
The chunks below are ranked by relevance to the question.

LOG ENTRIES:
{log_ctx}

KNOWLEDGE BASE ARTICLES:
{kb_ctx}

USER QUESTION: {req.q}

Provide a concise, accurate answer:
- Explain errors with specifics (timestamps, error codes, device IDs).
- Summarize relevant KB solutions.
- Reference specific sources.
- If insufficient info, say what's needed.

ANSWER:"""

    ptok = estimate_tokens(prompt)
    max_tok = TOKEN_BUDGET["MODEL_MAX_TOKENS"] - TOKEN_BUDGET["MAX_GENERATION_TOKENS"]
    logger.info("ASK: q='%s' | log=%d kb=%d | ~%d tok (limit %d)",
                req.q[:50], len(log_ctx), len(kb_ctx), ptok, max_tok)

    # Layer 3: safe_generate
    answer = safe_generate(prompt)
    ms = int((time.perf_counter() - start) * 1000)

    return AnswerResponse(
        answer=answer, log_sources=log_src, kb_sources=kb_src,
        confidence=confidence, processing_time_ms=ms,
        context_stats={
            "search_mode": "hybrid (vector + BM25 + re-ranking)",
            "log_candidates": len(log_cands), "kb_candidates": len(kb_cands),
            "log_after_rerank": len(log_ranked), "kb_after_rerank": len(kb_ranked),
            "log_context_chars": len(log_ctx), "kb_context_chars": len(kb_ctx),
            "prompt_chars": len(prompt), "prompt_tokens": ptok,
            "max_tokens": max_tok, "headroom": max_tok - ptok,
        },
    )


# ============================================================================
# ERROR HANDLERS
# ============================================================================
@app.exception_handler(HTTPException)
async def http_err(request, exc):
    return JSONResponse(exc.status_code, {
        "error": exc.detail, "path": request.url.path,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })

@app.exception_handler(Exception)
async def general_err(request, exc):
    logger.exception("Unhandled: %s", exc)
    return JSONResponse(500, {
        "error": "Internal server error", "path": request.url.path,
        "timestamp": datetime.now(timezone.utc).isoformat(),
    })